import json
from io import BytesIO
from datetime import date, datetime, timedelta
from ..utils.time import now_ph

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from flask import Blueprint, flash, jsonify, redirect, render_template, request, send_file, url_for
from flask_login import current_user, login_required
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font
from sqlalchemy import func

from ..decorators import admin_required
from ..extensions import db
from ..models.activity_log import ActivityLog
from ..models.category import Category
from ..models.person import Person
from ..models.attendance import Attendance
from ..models.event import Event, EventDay

admin_bp = Blueprint('admin', __name__, url_prefix='/admin', template_folder='../templates')


def get_valid_categories():
    categories = Category.query.order_by(Category.name).all()
    names = [category.name for category in categories]
    if not names:
        names = ['Ladies', 'Men', 'Young People', 'Young Professionals']
    return names


def get_active_categories():
    return Category.query.filter_by(is_active=True).order_by(Category.name).all()


def build_admin_metrics():
    total_people = Person.query.count()
    registered = Person.query.filter_by(registration_status='Registered').count()
    not_registered = Person.query.filter_by(registration_status='Not Registered').count()
    members = Person.query.filter_by(person_type='Member').count()
    visitors = Person.query.filter_by(person_type='Visitor').count()
    attendance_date = now_ph().date()
    present = Attendance.query.filter_by(attendance_date=attendance_date, status='Present').count()
    absent = max(0, registered - present)
    progress_pct = int((registered / total_people) * 100) if total_people else 0
    categories = Category.query.order_by(Category.name).all()
    category_counts = [
        {
            'name': category.name,
            'count': Person.query.filter_by(category=category.name).count(),
            'active': category.is_active,
        }
        for category in categories
    ]

    return {
        'total_people': total_people,
        'registered': registered,
        'not_registered': not_registered,
        'members': members,
        'visitors': visitors,
        'present': present,
        'absent': absent,
        'progress_pct': progress_pct,
        'category_counts': category_counts,
    }


def get_active_category_names():
    categories = Category.query.filter_by(is_active=True).order_by(Category.name).all()
    return [category.name for category in categories]


def get_day_label(day):
    if not day:
        return 'Current Day'
    return f"Day {day.day_number} — {day.date.strftime('%B %d, %Y')}"


def get_active_event():
    active_event = Event.query.filter_by(status='Active').order_by(Event.created_at.desc()).first()
    if active_event:
        return active_event
    return Event.query.order_by(Event.created_at.desc()).first()


def get_active_event_day(event=None):
    event = event or get_active_event()
    if not event:
        return None
    if event.active_day_id:
        day = EventDay.query.get(event.active_day_id)
        if day and day.event_id == event.id:
            return day
    day = event.event_days.order_by(EventDay.day_number).first()
    if day:
        event.active_day_id = day.id
        db.session.commit()
    return day


def get_event_day_for_request(event=None, default_day_id=None):
    event = event or get_active_event()
    if not event:
        return None
    selected_day_id = request.args.get('day_id') or default_day_id
    if selected_day_id:
        day = EventDay.query.filter_by(id=selected_day_id, event_id=event.id).first()
        if day:
            return day
    return get_active_event_day(event)


def get_admin_attendance_rows(category='All categories', search='', status='All', event=None, selected_day=None):
    event = event or get_active_event()
    selected_day = selected_day or get_active_event_day(event)
    query = Person.query.filter_by(registration_status='Registered')

    if category and category != 'All categories':
        query = query.filter_by(category=category)
    if search:
        query = query.filter(func.lower(Person.name).contains(search.lower()))

    rows = []
    for person in query.order_by(Person.name).all():
        attendance = None
        if event and selected_day:
            attendance = Attendance.query.filter_by(
                event_id=event.id,
                attendance_day_id=selected_day.id,
                person_id=person.id,
            ).order_by(Attendance.created_at.desc()).first()
        attendance_status = attendance.status if attendance else 'Absent'
        attendance_time = attendance.attendance_time if attendance else None

        if status == 'Present' and attendance_status != 'Present':
            continue
        if status == 'Absent' and attendance_status != 'Absent':
            continue

        rows.append({
            'person_id': person.id,
            'name': person.name,
            'category': person.category,
            'person_type': person.person_type,
            'registration_status': person.registration_status,
            'attendance_status': attendance_status,
            'attendance_time': attendance_time.strftime('%I:%M %p') if attendance_time else '—',
        })

    return rows


def build_admin_attendance_summary(category='All categories', search='', status='All', event=None, selected_day=None):
    rows = get_admin_attendance_rows(category=category, search=search, status=status, event=event, selected_day=selected_day)
    total_registered = len(rows)
    present = sum(1 for row in rows if row['attendance_status'] == 'Present')
    absent = sum(1 for row in rows if row['attendance_status'] == 'Absent')
    return {
        'total_registered': total_registered,
        'present': present,
        'absent': absent,
    }


def set_paragraph_font(paragraph, name='Arial', size=Pt(11), bold=False):
    for run in paragraph.runs:
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:rFonts'), name)
        run.font.size = size
        run.font.bold = bold


def build_person_table(document, people):
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['No.', 'Name', 'Category', 'Person Type', 'Registration Status']
    for idx, text in enumerate(headers):
        cell = hdr_cells[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, person in enumerate(people, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = person.name
        row_cells[2].text = person.category
        row_cells[3].text = person.person_type
        row_cells[4].text = person.registration_status

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:rFonts'), 'Arial')
                    run.font.size = Pt(10)

    return table


def add_section(document, title, people):
    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:rFonts'), 'Arial')
    run.font.size = Pt(14)
    run.font.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph()
    build_person_table(document, people)
    document.add_paragraph()


def reconcile_event_days(event, start_date, end_date):
    if not event:
        return

    desired_days = []
    total_days = (end_date - start_date).days + 1
    for index in range(total_days):
        current_date = start_date + timedelta(days=index)
        desired_days.append((index + 1, current_date))

    desired_numbers = {day_number for day_number, _ in desired_days}
    desired_dates = {day_date for _, day_date in desired_days}

    existing_days = list(event.event_days.all())
    for day in existing_days:
        if day.day_number not in desired_numbers and day.date not in desired_dates:
            db.session.delete(day)

    existing_by_date = {day.date: day for day in event.event_days.all()}
    existing_by_number = {day.day_number: day for day in event.event_days.all()}

    for day_number, current_date in desired_days:
        day = existing_by_date.get(current_date) or existing_by_number.get(day_number)
        if day is None:
            db.session.add(EventDay(
                event_id=event.id,
                day_number=day_number,
                date=current_date,
                is_active=False,
            ))
            continue

        day.day_number = day_number
        day.date = current_date
        day.is_active = (event.active_day_id == day.id)

    if event.status == 'Active' and event.active_day_id is None:
        first_day = EventDay.query.filter_by(event_id=event.id).order_by(EventDay.day_number).first()
        if first_day:
            event.active_day_id = first_day.id

    for day in event.event_days.order_by(EventDay.day_number).all():
        day.is_active = (event.active_day_id == day.id)


def create_export_document(category, people_query):
    document = Document()
    document.styles['Normal'].font.name = 'Arial'
    document.styles['Normal'].font.size = Pt(10)

    title = document.add_paragraph()
    title_run = title.add_run('HOME BUILDERS')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run('REGISTRATION RECORD')
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if category != 'All':
        section_title = document.add_paragraph()
        section_run = section_title.add_run(category.upper())
        section_run.font.name = 'Arial'
        section_run.font.size = Pt(14)
        section_run.font.bold = True
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    date_paragraph = document.add_paragraph()
    date_run = date_paragraph.add_run(f'Date Generated: {date.today().strftime("%B %d, %Y")}')
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(10)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()

    if category == 'All':
        for section in ['Ladies', 'Men', 'Young People']:
            subset = people_query.filter_by(category=section).order_by(Person.name).all()
            if subset:
                add_section(document, section.upper(), subset)
    else:
        people = people_query.order_by(Person.name).all()
        add_section(document, category.upper(), people)

    document.add_paragraph()
    document.add_paragraph('SUMMARY').runs[0].font.bold = True
    document.add_paragraph()

    stats = {
        'members_registered': people_query.filter_by(person_type='Member', registration_status='Registered').count(),
        'members_not_registered': people_query.filter_by(person_type='Member', registration_status='Not Registered').count(),
        'visitors_registered': people_query.filter_by(person_type='Visitor', registration_status='Registered').count(),
        'visitors_not_registered': people_query.filter_by(person_type='Visitor', registration_status='Not Registered').count(),
        'total': people_query.count(),
        'registered': people_query.filter_by(registration_status='Registered').count(),
        'not_registered': people_query.filter_by(registration_status='Not Registered').count(),
    }

    summary_table = document.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.cell(0, 0).text = 'Members:'
    summary_table.cell(0, 1).text = ''
    summary_table.cell(1, 0).text = '  Registered'
    summary_table.cell(1, 1).text = str(stats['members_registered'])
    summary_table.cell(2, 0).text = '  Not Registered'
    summary_table.cell(2, 1).text = str(stats['members_not_registered'])
    summary_table.cell(3, 0).text = 'Visitors:'
    summary_table.cell(3, 1).text = ''

    summary_table.add_row().cells[0].text = '  Registered'
    summary_table.add_row().cells[1].text = str(stats['visitors_registered'])
    summary_table.add_row().cells[0].text = '  Not Registered'
    summary_table.add_row().cells[1].text = str(stats['visitors_not_registered'])

    document.add_paragraph()
    total_table = document.add_table(rows=3, cols=2)
    total_table.style = 'Table Grid'
    total_table.cell(0, 0).text = 'Total'
    total_table.cell(0, 1).text = str(stats['total'])
    total_table.cell(1, 0).text = 'Registered'
    total_table.cell(1, 1).text = str(stats['registered'])
    total_table.cell(2, 0).text = 'Not Registered'
    total_table.cell(2, 1).text = str(stats['not_registered'])

    return document


VALID_CATEGORIES = ['Ladies', 'Men', 'Young People']
VALID_PERSON_TYPES = ['Member', 'Visitor']
VALID_REGISTRATION_STATUSES = ['Registered', 'Not Registered']
REQUIRED_IMPORT_HEADERS = ['Name', 'Category', 'Person Type', 'Registration Status']


def create_export_workbook(category, people_query):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = category if category != 'All' else 'People'
    headers = ['Name', 'Category', 'Person Type', 'Registration Status', 'Registered At']
    sheet.append(headers)

    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')

    for person in people_query.order_by(Person.name):
        sheet.append([
            person.name,
            person.category,
            person.person_type,
            person.registration_status,
            person.registered_at.strftime('%Y-%m-%d %H:%M') if person.registered_at else '',
        ])

    for column_cells in sheet.columns:
        max_length = max((len(str(cell.value)) if cell.value is not None else 0) for cell in column_cells)
        width = min(max_length + 2, 40)
        sheet.column_dimensions[column_cells[0].column_letter].width = width

    return workbook


def normalize_name(value):
    return ' '.join(str(value or '').strip().split()).lower()


def parse_import_file(file_storage):
    filename = file_storage.filename or ''
    if not filename.lower().endswith('.xlsx'):
        return [], ['Only .xlsx files are supported for import.']

    try:
        workbook = load_workbook(filename=file_storage, data_only=True)
    except Exception:
        return [], ['Unable to read the Excel file. Please upload a valid .xlsx workbook.']

    sheet = workbook.active
    header_row = [str(cell.value).strip() if cell.value else '' for cell in next(sheet.iter_rows(min_row=1, max_row=1))]
    header_map = {header.lower(): header for header in header_row if header}
    missing_headers = [required for required in REQUIRED_IMPORT_HEADERS if required.lower() not in header_map]
    if missing_headers:
        return [], [f'Missing required header: {header}' for header in missing_headers]

    rows = []
    for row_idx, row in enumerate(sheet.iter_rows(min_row=2), start=2):
        if all(cell.value is None or str(cell.value).strip() == '' for cell in row):
            continue

        row_data = {}
        for idx, cell in enumerate(row):
            if idx >= len(header_row):
                continue
            header = header_row[idx]
            normalized_header = header.lower().strip()
            if normalized_header in header_map:
                row_data[header_map[normalized_header]] = str(cell.value).strip() if cell.value is not None else ''

        rows.append(row_data)

    if not rows:
        return [], ['The file does not contain any data rows.']

    return rows, []


def validate_import_rows(rows):
    errors = []
    warnings = []
    valid_rows = []
    seen_names = set()
    existing_names = {normalize_name(person.name): person for person in Person.query.all()}

    for index, row in enumerate(rows, start=2):
        row_errors = []
        name = row.get('Name', '').strip()
        category = row.get('Category', '').strip()
        person_type = row.get('Person Type', '').strip()
        registration_status = row.get('Registration Status', '').strip()
        normalized_name = normalize_name(name)

        if not name:
            row_errors.append('Name is required.')
        if category not in VALID_CATEGORIES:
            row_errors.append(f"Category must be one of {', '.join(VALID_CATEGORIES)}.")
        if person_type not in VALID_PERSON_TYPES:
            row_errors.append(f"Person Type must be one of {', '.join(VALID_PERSON_TYPES)}.")
        if registration_status not in VALID_REGISTRATION_STATUSES:
            row_errors.append(f"Registration Status must be 'Registered' or 'Not Registered'.")

        if normalized_name:
            if normalized_name in seen_names:
                row_errors.append('Duplicate name found in uploaded data.')
            else:
                seen_names.add(normalized_name)

            if normalized_name in existing_names:
                warnings.append(f"Row {index}: '{name}' looks like an existing person already in the database.")

        if row_errors:
            errors.append(f'Row {index}: ' + ' '.join(row_errors))
        else:
            valid_rows.append({
                'Name': name,
                'Category': category,
                'Person Type': person_type,
                'Registration Status': registration_status,
            })

    return valid_rows, errors, warnings


def import_people_from_rows(rows, user):
    imported_count = 0
    existing_names = {normalize_name(person.name) for person in Person.query.all()}

    for row in rows:
        normalized_name = normalize_name(row['Name'])
        if not row['Name'] or normalized_name in existing_names:
            continue

        person = Person(
            name=row['Name'],
            category=row['Category'],
            person_type=row['Person Type'],
            registration_status=row['Registration Status'],
        )
        if person.registration_status == 'Registered':
            person.registered_at = now_ph()

        db.session.add(person)
        db.session.flush()

        activity = ActivityLog(
            user_id=user.id,
            person_id=person.id,
            action='Import person',
            description=f'Imported {person.name} via Excel upload',
        )
        db.session.add(activity)

        existing_names.add(normalized_name)
        imported_count += 1

    if imported_count > 0:
        db.session.commit()

    return imported_count


def apply_filters(query, category, person_type, registration_status):
    if category and category != 'All':
        query = query.filter_by(category=category)
    if person_type and person_type != 'All':
        query = query.filter_by(person_type=person_type)
    if registration_status and registration_status != 'All':
        query = query.filter_by(registration_status=registration_status)
    return query


@admin_bp.route('/')
@login_required
@admin_required
def dashboard():
    metrics = build_admin_metrics()
    activities = ActivityLog.query.order_by(ActivityLog.created_at.desc()).limit(15).all()
    return render_template('admin/dashboard.html', activities=activities, **metrics)


@admin_bp.route('/attendance')
@login_required
@admin_required
def attendance():
    category = request.args.get('category', 'All categories').strip() or 'All categories'
    search = request.args.get('search', '').strip()
    status = request.args.get('status', 'All').strip() or 'All'
    event = get_active_event()
    selected_day = get_event_day_for_request(event=event)
    summary = build_admin_attendance_summary(category=category, search=search, status=status, event=event, selected_day=selected_day)
    rows = get_admin_attendance_rows(category=category, search=search, status=status, event=event, selected_day=selected_day)
    categories = ['All categories'] + get_active_category_names()
    event_days = event.event_days.order_by(EventDay.day_number).all() if event else []
    return render_template(
        'admin/attendance.html',
        summary=summary,
        rows=rows,
        categories=categories,
        selected_category=category,
        selected_status=status,
        search=search,
        current_event=event,
        current_day=selected_day,
        event_days=event_days,
        view_all=request.args.get('view_all', '0') == '1',
    )


@admin_bp.route('/attendance/data')
@login_required
@admin_required
def attendance_data():
    category = request.args.get('category', 'All categories').strip() or 'All categories'
    search = request.args.get('search', '').strip()
    status = request.args.get('status', 'All').strip() or 'All'
    event = get_active_event()
    selected_day = get_event_day_for_request(event=event)
    summary = build_admin_attendance_summary(category=category, search=search, status=status, event=event, selected_day=selected_day)
    rows = get_admin_attendance_rows(category=category, search=search, status=status, event=event, selected_day=selected_day)
    category_names = ['All categories'] + get_active_category_names()
    return jsonify({
        'summary': summary,
        'rows': rows,
        'categories': category_names,
        'event_name': event.name if event else '',
        'day_label': f"Day {selected_day.day_number} — {selected_day.date.strftime('%B %d, %Y')}" if selected_day else '',
    })


@admin_bp.route('/attendance/update-status', methods=['POST'])
@login_required
@admin_required
def update_attendance_status():
    person_id = request.form.get('person_id')
    status = request.form.get('status', 'Present').strip()
    if not person_id or status not in {'Present', 'Absent'}:
        return jsonify({'success': False, 'message': 'Invalid attendance update.'}), 400

    person = Person.query.get_or_404(person_id)
    if person.registration_status != 'Registered':
        return jsonify({'success': False, 'message': 'Only registered people can be marked for attendance.'}), 400

    event = get_active_event()
    selected_day = get_event_day_for_request(event=event)
    if not event or not selected_day:
        return jsonify({'success': False, 'message': 'No active event or attendance day is available.'}), 400

    now_time = now_ph().time()
    record = Attendance.query.filter_by(
        event_id=event.id,
        attendance_day_id=selected_day.id,
        person_id=person.id,
    ).first()

    if status == 'Present':
        if record is None:
            record = Attendance(
                event_id=event.id,
                attendance_day_id=selected_day.id,
                person_id=person.id,
                attendance_date=selected_day.date,
                attendance_time=now_time,
                status='Present',
                created_by=current_user.id,
            )
            db.session.add(record)
        else:
            record.attendance_time = now_time
            record.status = 'Present'
            record.created_by = current_user.id
    else:
        if record is None:
            record = Attendance(
                event_id=event.id,
                attendance_day_id=selected_day.id,
                person_id=person.id,
                attendance_date=selected_day.date,
                attendance_time=now_time,
                status='Absent',
                created_by=current_user.id,
            )
            db.session.add(record)
        else:
            record.attendance_time = now_time
            record.status = 'Absent'
            record.created_by = current_user.id

    db.session.commit()

    activity = ActivityLog(
        user_id=current_user.id,
        person_id=person.id,
        action='Update attendance status',
        description=f'Admin updated {person.name} attendance to {status} for {event.name} Day {selected_day.day_number}.',
    )
    db.session.add(activity)
    db.session.commit()

    socketio = __import__('app.extensions', fromlist=['socketio']).socketio
    socketio.emit('attendance_update', {
        'name': person.name,
        'category': person.category,
        'person_type': person.person_type,
        'status': status,
        'time': now_time.strftime('%I:%M %p'),
        'event': event.name,
        'day': f"Day {selected_day.day_number} — {selected_day.date.strftime('%B %d, %Y')}",
    })

    return jsonify({'success': True, 'status': status, 'attendance_time': now_time.strftime('%I:%M %p')})


@admin_bp.route('/attendance/export')
@login_required
@admin_required
def export_attendance():
    export_format = request.args.get('format', 'xlsx').lower()
    category = request.args.get('category', 'All categories').strip() or 'All categories'
    status = request.args.get('status', 'All').strip() or 'All'
    search = request.args.get('search', '').strip()
    export_all = str(request.args.get('all_results', 'false')).lower() in {'1', 'true', 'yes'}
    event = get_active_event()
    selected_day = get_event_day_for_request(event=event)
    selected_category = 'All categories' if export_all else category
    selected_status = 'All' if export_all else status

    if export_all and event:
        event_days = event.event_days.order_by(EventDay.day_number).all()
        day_rows = []
        for day in event_days:
            day_rows.extend(get_admin_attendance_rows(category='All categories', search='', status='All', event=event, selected_day=day))
        rows = day_rows
    else:
        rows = get_admin_attendance_rows(category=selected_category, search=search, status=selected_status, event=event, selected_day=selected_day)

    activity = ActivityLog(
        user_id=current_user.id,
        action='Export attendance',
        description=f'Admin exported {event.name if event else "Attendance"} {get_day_label(selected_day) if selected_day else "report"} {export_format.upper()} report.',
    )
    db.session.add(activity)
    db.session.commit()

    if export_format == 'xlsx':
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = 'Attendance'
        sheet.append(['HOME BUILDERS ATTENDANCE REPORT'])
        sheet.append([])
        sheet.append(['Event:', event.name if event else ''])
        sheet.append(['Day:', f"Day {selected_day.day_number} — {selected_day.date.strftime('%B %d, %Y')}" if selected_day else 'All Days'])
        sheet.append(['Date:', now_ph().strftime('%B %d, %Y')])
        sheet.append([])
        headers = ['No.', 'Name', 'Category', 'Person Type', 'Registration Status', 'Attendance', 'Time']
        sheet.append(headers)

        for cell in sheet[1]:
            cell.font = Font(bold=True, size=12)

        for idx, row in enumerate(rows, start=1):
            sheet.append([
                idx,
                row['name'],
                row['category'],
                row['person_type'],
                row['registration_status'],
                row['attendance_status'],
                row['attendance_time'],
            ])

        for column_cells in sheet.columns:
            max_length = max((len(str(cell.value)) if cell.value is not None else 0) for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = min(max_length + 2, 28)

        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        filename = f"home_builders_attendance_{(event.name if event else 'event').lower().replace(' ', '_')}_{now_ph().strftime('%Y%m%d')}.xlsx"
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    document = Document()
    document.add_heading('HOME BUILDERS ATTENDANCE REPORT', level=1)
    event_label = event.name if event else ''
    day_label = 'All Days' if export_all else get_day_label(selected_day) if selected_day else ''
    document.add_paragraph(f'Event: {event_label}')
    document.add_paragraph(f'Day: {day_label}')
    document.add_paragraph(f'Date: {now_ph().strftime("%B %d, %Y")}')
    document.add_paragraph('Summary')
    summary = build_admin_attendance_summary(category=selected_category, search=search if not export_all else '', status=selected_status if not export_all else 'All', event=event, selected_day=selected_day)
    stats_table = document.add_table(rows=1, cols=2)
    stats_table.style = 'Table Grid'
    hdr = stats_table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    for key, label in [('total_registered', 'Total Registered'), ('present', 'Present'), ('absent', 'Absent')]:
        row = stats_table.add_row().cells
        row[0].text = label
        row[1].text = str(summary[key])

    grouped = {}
    if export_all and event:
        for day in event.event_days.order_by(EventDay.day_number).all():
            day_rows = get_admin_attendance_rows(category='All categories', search='', status='All', event=event, selected_day=day)
            grouped[f"Day {day.day_number} — {day.date.strftime('%B %d, %Y')}"] = day_rows
    else:
        grouped[f"Day {selected_day.day_number} — {selected_day.date.strftime('%B %d, %Y')}" if selected_day else 'Current Report'] = rows

    for label, category_rows in grouped.items():
        document.add_paragraph(label.upper())
        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        headers = ['No.', 'Name', 'Category', 'Person Type', 'Registration', 'Attendance', 'Time']
        header_row = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_row[idx].text = header
        for idx, row in enumerate(category_rows, start=1):
            values = [
                str(idx),
                row['name'],
                row['category'],
                row['person_type'],
                row['registration_status'],
                row['attendance_status'],
                row['attendance_time'],
            ]
            data_cells = table.add_row().cells
            for column_idx, value in enumerate(values):
                data_cells[column_idx].text = str(value)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f"home_builders_attendance_{(event.name if event else 'event').lower().replace(' ', '_')}_{now_ph().strftime('%Y%m%d')}.docx"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@admin_bp.route('/events')
@login_required
@admin_required
def events():
    events = Event.query.order_by(Event.start_date.asc(), Event.created_at.asc()).all()
    for event in events:
        event.days = event.event_days.order_by(EventDay.day_number).all()
    active_event = get_active_event()
    return render_template('admin/events.html', events=events, active_event=active_event)


@admin_bp.route('/events/create', methods=['POST'])
@login_required
@admin_required
def create_event():
    name = (request.form.get('name') or '').strip()
    start_date_str = (request.form.get('start_date') or '').strip()
    end_date_str = (request.form.get('end_date') or '').strip()
    status = (request.form.get('status') or 'Upcoming').strip()

    if not name or not start_date_str or not end_date_str:
        flash('Event name and dates are required.', 'danger')
        return redirect(url_for('admin.events'))

    try:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
    except ValueError:
        flash('Use valid dates in YYYY-MM-DD format.', 'danger')
        return redirect(url_for('admin.events'))

    if end_date < start_date:
        flash('Event end date must be on or after the start date.', 'danger')
        return redirect(url_for('admin.events'))

    existing = Event.query.filter(db.func.lower(Event.name) == name.lower()).first()
    if existing:
        flash('An event with that name already exists.', 'warning')
        return redirect(url_for('admin.events'))

    event = Event(
        name=name,
        start_date=start_date,
        end_date=end_date,
        status=status if status in {'Upcoming', 'Active', 'Completed'} else 'Upcoming',
    )
    db.session.add(event)
    db.session.flush()

    reconcile_event_days(event, start_date, end_date)
    if event.status == 'Active' and event.active_day_id is None:
        first_day = EventDay.query.filter_by(event_id=event.id).order_by(EventDay.day_number).first()
        if first_day:
            event.active_day_id = first_day.id

    db.session.commit()
    flash(f'Event {event.name} created successfully.', 'success')
    return redirect(url_for('admin.events'))


@admin_bp.route('/events/<int:event_id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_event(event_id):
    event = Event.query.get_or_404(event_id)
    name = (request.form.get('name') or '').strip()
    start_date_str = (request.form.get('start_date') or '').strip()
    end_date_str = (request.form.get('end_date') or '').strip()
    status = (request.form.get('status') or event.status).strip()

    if not name or not start_date_str or not end_date_str:
        flash('Event name and dates are required.', 'danger')
        return redirect(url_for('admin.events'))

    try:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
    except ValueError:
        flash('Use valid dates in YYYY-MM-DD format.', 'danger')
        return redirect(url_for('admin.events'))

    if end_date < start_date:
        flash('Event end date must be on or after the start date.', 'danger')
        return redirect(url_for('admin.events'))

    event.name = name
    event.start_date = start_date
    event.end_date = end_date
    event.status = status if status in {'Upcoming', 'Active', 'Completed'} else event.status
    db.session.flush()

    reconcile_event_days(event, start_date, end_date)
    if event.status == 'Active' and event.active_day_id is None:
        first_day = EventDay.query.filter_by(event_id=event.id).order_by(EventDay.day_number).first()
        if first_day:
            event.active_day_id = first_day.id

    db.session.commit()
    flash('Event updated successfully.', 'success')
    return redirect(url_for('admin.events'))


@admin_bp.route('/events/<int:event_id>/set-active-day', methods=['POST'])
@login_required
@admin_required
def set_active_day(event_id):
    event = Event.query.get_or_404(event_id)
    day_id = request.form.get('day_id')
    if not day_id:
        flash('Please select an attendance day.', 'danger')
        return redirect(url_for('admin.events'))

    day = EventDay.query.filter_by(id=day_id, event_id=event.id).first()
    if day is None:
        flash('Selected attendance day is invalid.', 'danger')
        return redirect(url_for('admin.events'))

    event.active_day_id = day.id
    event.status = 'Active'
    db.session.commit()
    flash(f'Active attendance day set to Day {day.day_number}.', 'success')
    return redirect(url_for('admin.events'))


@admin_bp.route('/events/<int:event_id>/status', methods=['POST'])
@login_required
@admin_required
def set_event_status(event_id):
    event = Event.query.get_or_404(event_id)
    status = (request.form.get('status') or event.status).strip()
    if status not in {'Upcoming', 'Active', 'Completed'}:
        status = event.status
    event.status = status
    db.session.commit()
    flash(f'Event status updated to {status}.', 'success')
    return redirect(url_for('admin.events'))


@admin_bp.route('/attendance/clear-filters')
@login_required
@admin_required
def clear_attendance_filters():
    return redirect(url_for('admin.attendance'))


@admin_bp.route('/categories')
@login_required
@admin_required
def categories():
    categories = Category.query.order_by(Category.name).all()
    return render_template('admin/categories.html', categories=categories)


@admin_bp.route('/categories/add', methods=['POST'])
@login_required
@admin_required
def add_category():
    name = (request.form.get('name') or '').strip()
    if not name:
        flash('Category name is required.', 'error')
        return redirect(url_for('admin.categories'))

    normalized = name.lower()
    if Category.query.filter(db.func.lower(Category.name) == normalized).first():
        flash('A category with that name already exists.', 'error')
        return redirect(url_for('admin.categories'))

    category = Category(name=name, is_active=True)
    db.session.add(category)
    db.session.commit()
    flash('Category added successfully.', 'success')
    return redirect(url_for('admin.categories'))


@admin_bp.route('/categories/edit', methods=['POST'])
@login_required
@admin_required
def edit_category():
    category_id = request.form.get('category_id')
    name = (request.form.get('name') or '').strip()
    if not category_id or not name:
        flash('Category information is missing.', 'error')
        return redirect(url_for('admin.categories'))

    category = Category.query.get(category_id)
    if not category:
        flash('Category not found.', 'error')
        return redirect(url_for('admin.categories'))

    if Category.query.filter(Category.id != category.id, db.func.lower(Category.name) == name.lower()).first():
        flash('A category with that name already exists.', 'error')
        return redirect(url_for('admin.categories'))

    category.name = name
    db.session.commit()
    flash('Category updated successfully.', 'success')
    return redirect(url_for('admin.categories'))


@admin_bp.route('/categories/<int:category_id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_category(category_id):
    category = Category.query.get_or_404(category_id)
    category.is_active = not category.is_active
    db.session.commit()
    flash(f"Category '{category.name}' was {'activated' if category.is_active else 'deactivated'}.", 'success')
    return redirect(url_for('admin.categories'))


@admin_bp.route('/export', methods=['GET', 'POST'])
@login_required
@admin_required
def export_page():
    categories = ['All', 'Ladies', 'Men', 'Young People']
    person_types = ['All', 'Member', 'Visitor']
    statuses = ['All', 'Registered', 'Not Registered']
    formats = [('docx', 'Word (.docx)'), ('xlsx', 'Excel (.xlsx)')]

    if request.method == 'POST':
        category = request.form.get('category', 'All')
        person_type = request.form.get('person_type', 'All')
        registration_status = request.form.get('registration_status', 'All')
        export_format = request.form.get('export_format', 'docx')

        query = Person.query
        filtered_query = apply_filters(query, category, person_type, registration_status)

        activity = ActivityLog(
            user_id=current_user.id,
            action='Export generated',
            description=f'Admin exported {category} {export_format.upper()} registration report',
        )
        db.session.add(activity)
        db.session.commit()

        if export_format == 'xlsx':
            workbook = create_export_workbook(category, filtered_query)
            buffer = BytesIO()
            workbook.save(buffer)
            buffer.seek(0)
            filename = f"home_builders_export_{category.lower().replace(' ', '_')}_{date.today().strftime('%Y%m%d')}.xlsx"
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            )

        document = create_export_document(category, filtered_query)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = f"home_builders_export_{category.lower().replace(' ', '_')}_{date.today().strftime('%Y%m%d')}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

    return render_template(
        'admin/export.html',
        categories=categories,
        person_types=person_types,
        statuses=statuses,
        formats=formats,
    )


@admin_bp.route('/import', methods=['GET', 'POST'])
@login_required
@admin_required
def import_page():
    preview_rows = []
    import_errors = []
    import_warnings = []
    imported_count = None
    import_ready = False

    if request.method == 'POST':
        if 'preview' in request.form:
            upload = request.files.get('import_file')
            if not upload or upload.filename == '':
                import_errors.append('Please select an Excel (.xlsx) file to preview.')
            else:
                rows, parse_errors = parse_import_file(upload)
                if parse_errors:
                    import_errors.extend(parse_errors)
                else:
                    preview_rows, validation_errors, validation_warnings = validate_import_rows(rows)
                    import_errors.extend(validation_errors)
                    import_warnings.extend(validation_warnings)
                    import_ready = bool(preview_rows) and not validation_errors and not validation_warnings

        elif 'confirm_import' in request.form:
            import_data = request.form.get('import_data')
            if not import_data:
                import_errors.append('No preview data available. Please upload and preview the file again.')
            else:
                try:
                    parsed_rows = json.loads(import_data)
                    imported_count = import_people_from_rows(parsed_rows, current_user)
                    if imported_count:
                        activity = ActivityLog(
                            user_id=current_user.id,
                            action='Import completed',
                            description=f'Admin imported {imported_count} records from Excel upload',
                        )
                        db.session.add(activity)
                        db.session.commit()
                        import_warnings = []
                        preview_rows = []
                    else:
                        import_errors.append('No new people were imported. The uploaded rows may already exist.')
                except json.JSONDecodeError:
                    import_errors.append('Unable to read preview data. Please retry the upload and preview step.')

    return render_template(
        'admin/import.html',
        preview_rows=preview_rows,
        import_errors=import_errors,
        import_warnings=import_warnings,
        imported_count=imported_count,
        import_ready=import_ready,
    )


@admin_bp.route('/activity')
@login_required
@admin_required
def activity():
    activities = ActivityLog.query.order_by(ActivityLog.created_at.desc()).limit(30).all()
    return render_template('admin/activity.html', activities=activities)


@admin_bp.route('/activity/clear', methods=['POST'])
@login_required
@admin_required
def clear_activity_history():
    try:
        deleted = ActivityLog.query.delete()
        db.session.commit()
        return jsonify({
            'success': True,
            'deleted_count': deleted,
            'message': 'Activity history cleared successfully.'
        })
    except Exception:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': 'Unable to clear activity history.'
        }), 500


@admin_bp.route('/metrics/json')
@login_required
@admin_required
def metrics_json():
    return jsonify(build_admin_metrics())


@admin_bp.route('/activity/json')
@login_required
@admin_required
def activity_json():
    activities = ActivityLog.query.order_by(ActivityLog.created_at.desc()).limit(30).all()
    payload = []
    for activity in activities:
        payload.append({
            'timestamp': activity.created_at.strftime('%I:%M %p'),
            'user': activity.user.username if activity.user else 'Unknown',
            'action': activity.action,
            'description': activity.description,
        })
    return jsonify(payload)
